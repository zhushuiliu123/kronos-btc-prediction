#!/usr/bin/env python3
"""
Convert Kronos optimization plan to a Word document (.docx)
"""

import os

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    import subprocess
    subprocess.check_call([
        r"C:\Users\DCU459-41\.workbuddy\binaries\python\versions\3.13.12\python.exe",
        "-m", "pip", "install", "python-docx", "-q"
    ])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table(doc, headers, rows, col_widths=None, header_color="4472C4"):
    """Add a styled table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph("")
    return table


def create_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== Title =====
    title = doc.add_heading('Kronos \u6a21\u578b\u4f18\u5316\u8ba1\u5212', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('\u57fa\u4e8e\u5b9e\u6d4b\u7ed3\u679c\u7684\u53ef\u6267\u884c\u4f18\u5316\u65b9\u6848')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    run.italic = True

    doc.add_paragraph('')

    # ===== Section 1: Problem Diagnosis =====
    doc.add_heading('\u4e00\u3001\u95ee\u9898\u8bca\u65ad\uff08\u57fa\u4e8e\u5b9e\u6d4b\u7ed3\u679c\uff09', level=1)

    p = doc.add_paragraph()
    run = p.add_run('\u57fa\u4e8e\u5bf9\u6bd4\u7279\u5e01 500 \u5929\u5386\u53f2\u6570\u636e\uff0cKronos-mini \u6a21\u578b\u9884\u6d4b 30 \u5929\u7684\u8bef\u5dee\u5206\u6790\u7ed3\u679c\u5982\u4e0b\uff1a')
    run.font.size = Pt(11)

    add_table(doc,
        ['\u6307\u6807', '\u5b9e\u6d4b\u503c', '\u8bc4\u4ef7'],
        [
            ['\u7cfb\u7edf\u6027\u504f\u9ad8\u9884\u6d4b', '29/30 \u5929\u9884\u6d4b\u503c > \u5b9e\u9645\u503c', '\u4e25\u91cd\u7cfb\u7edf\u504f\u5dee'],
            ['\u5e73\u5747\u7edd\u5bf9\u8bef\u5dee\uff08MAE\uff09', '$21,104', '\u504f\u9ad8'],
            ['\u5e73\u5747\u767e\u5206\u6bd4\u8bef\u5dee\uff08MAPE\uff09', '30.17%', '\u504f\u9ad8'],
            ['\u5747\u65b9\u6839\u8bef\u5dee\uff08RMSE\uff09', '$24,425', '\u504f\u9ad8'],
            ['\u6da8\u8dcc\u65b9\u5411\u51c6\u786e\u7387', '69%', '\u5c1a\u53ef\uff0c\u6bd4\u968f\u673a\u597d'],
            ['\u6839\u672c\u95ee\u9898', '\u6a21\u578b\u8f93\u51fa\u5206\u5e03\u504f\u79fb\uff08bias toward higher values\uff09', '\u9700\u8981\u4fee\u590d'],
        ]
    )

    # ===== Section 2: Optimization Plans =====
    doc.add_heading('\u4e8c\u3001\u4f18\u5316\u65b9\u6848\uff08\u6309\u4f18\u5148\u7ea7\u6392\u5e8f\uff09', level=1)

    # --- Plan 1 ---
    doc.add_heading('\u65b9\u6848\u4e00\uff1a\u63a8\u7406\u53c2\u6570\u8c03\u4f18', level=2)
    p = doc.add_paragraph()
    run = p.add_run('\u3010\u6700\u4f18\u5148\uff0c\u7acb\u5373\u53ef\u505a\u3011')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

    p = doc.add_paragraph()
    run = p.add_run('\u95ee\u9898\uff1a')
    run.bold = True
    p.add_run('Temperature=1.0 \u5bfc\u81f4\u9884\u6d4b\u53d1\u6563\uff0c\u91c7\u6837\u504f\u5411\u9ad8\u503c\u533a\u95f4')

    p = doc.add_paragraph()
    run = p.add_run('\u884c\u52a8\u8ba1\u5212\uff1a')
    run.bold = True
    doc.add_paragraph('\u5728\u5de5\u4f5c\u53f0\u6d4b\u8bd5\u4e0d\u540c Temperature \u53c2\u6570\u7ec4\u5408\uff08\u6279\u91cf\u9884\u6d4b\u529f\u80fd\uff09')
    doc.add_paragraph('T=0.3 / T=0.5 / T=0.7 \u5404\u8dd1 5 \u6b21\uff0c\u53d6\u5e73\u5747\uff0c\u56fa\u5b9a top_p=0.9', style='List Bullet')
    doc.add_paragraph('\u52a0\u5165\u540e\u6821\u51c6\uff08Post-hoc Calibration\uff09\uff1a\u7528\u5386\u53f2\u6570\u636e\u8ba1\u7b97\u504f\u5dee\u7cfb\u6570\uff0c\u5bf9\u9884\u6d4b\u7ed3\u679c\u505a\u7f29\u653e\u6821\u6b63', style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('\u9884\u671f\u6548\u679c\uff1a')
    run.bold = True
    run = p.add_run('MAPE \u4ece 30% \u2192 15~20%')
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

    # --- Plan 2 ---
    doc.add_heading('\u65b9\u6848\u4e8c\uff1a\u6362\u66f4\u5927\u6a21\u578b', level=2)
    p = doc.add_paragraph()
    run = p.add_run('\u3010\u63a8\u8350\uff0c\u6548\u679c\u6700\u660e\u663e\u3011')
    run.bold = True
    run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)

    p = doc.add_paragraph()
    run = p.add_run('\u95ee\u9898\uff1a')
    run.bold = True
    p.add_run('mini \u7248\uff0870M \u53c2\u6570\uff09\u5bb9\u91cf\u592a\u5c0f\uff0c\u62df\u5408\u80fd\u529b\u4e0d\u8db3')

    add_table(doc,
        ['\u6a21\u578b', '\u53c2\u6570\u91cf', '\u4e0a\u4e0b\u6587', '\u9884\u8ba1\u6548\u679c'],
        [
            ['Kronos-mini\uff08\u5f53\u524d\uff09', '~70M', '1024', '\u57fa\u7ebf'],
            ['Kronos-small', '~180M', '2048', 'MAPE \u964d\u81f3 15~20%'],
            ['Kronos-base', '~400M', '2048', 'MAPE \u964d\u81f3 10~15%'],
            ['Kronos-large', '~1B', '4096', '\u6700\u4f73\u6548\u679c'],
        ]
    )

    p = doc.add_paragraph()
    run = p.add_run('\u884c\u52a8\u8ba1\u5212\uff1a')
    run.bold = True
    doc.add_paragraph('\u4e0b\u8f7d Kronos-small \u6743\u91cd\uff08HuggingFace: NeoQuasar/Kronos-small\uff09', style='List Bullet')
    doc.add_paragraph('\u4fee\u6539 kronos_numpy/ \u652f\u6301\u52a0\u8f7d\u4e0d\u540c\u5c3a\u5bf8\u6a21\u578b', style='List Bullet')
    doc.add_paragraph('\u7528\u76f8\u540c\u6570\u636e\u5bf9\u6bd4 mini vs small vs base \u7684 MAPE', style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('\u6210\u672c\uff1a')
    run.bold = True
    p.add_run('\u4e0b\u8f7d ~2GB\uff08small\uff09\u5230 ~8GB\uff08base\uff09')

    # --- Plan 3 ---
    doc.add_heading('\u65b9\u6848\u4e09\uff1a\u6ed1\u52a8\u7a97\u53e3\u56de\u6eaf\u6d4b\u8bd5', level=2)
    p = doc.add_paragraph()
    run = p.add_run('\u95ee\u9898\uff1a')
    run.bold = True
    p.add_run('\u5f53\u524d\u53ea\u6d4b\u4e86 30 \u5929\uff0c\u6837\u672c\u592a\u5c11\uff0c\u7ed3\u8bba\u4e0d\u53ef\u9760')

    p = doc.add_paragraph()
    run = p.add_run('\u884c\u52a8\u8ba1\u5212\uff1a')
    run.bold = True
    doc.add_paragraph('\u7528 500 \u5929\u6570\u636e\u505a\u6eda\u52a8\u9884\u6d4b\u6d4b\u8bd5\uff0c\u5171\u6eda\u52a8 200 \u4e2a\u7a97\u53e3', style='List Bullet')
    doc.add_paragraph('\u8f93\u51fa\u7efc\u5408\u6307\u6807\uff1a\u6574\u4f53 MAE / MAPE / \u65b9\u5411\u51c6\u786e\u7387', style='List Bullet')
    doc.add_paragraph('\u5206\u6790\u4e0d\u540c\u5e02\u573a\u72b6\u6001\u4e0b\u7684\u8868\u73b0\uff08\u725b\u5e02 vs \u718a\u5e02 vs \u9707\u8361\uff09', style='List Bullet')

    # --- Plan 4 ---
    doc.add_heading('\u65b9\u6848\u56db\uff1a\u7279\u5f81\u5de5\u7a0b\u4f18\u5316', level=2)
    p = doc.add_paragraph()
    run = p.add_run('\u95ee\u9898\uff1a')
    run.bold = True
    p.add_run('\u5f53\u524d\u53ea\u7528 OHLCV 5 \u4e2a\u7279\u5f81\uff0c\u4fe1\u606f\u91cf\u4e0d\u8db3')

    add_table(doc,
        ['\u7279\u5f81', '\u8ba1\u7b97\u65b9\u5f0f', '\u4f5c\u7528'],
        [
            ['RSI(14)', '\u76f8\u5bf9\u5f3a\u5f31\u6307\u6807', '\u6355\u6349\u8d85\u4e70\u8d85\u5356'],
            ['SMA(20)', '20\u65e5\u5747\u7ebf', '\u8d8b\u52bf\u65b9\u5411'],
            ['ATR(14)', '\u771f\u5b9e\u6ce2\u5e45', '\u6ce2\u52a8\u7387'],
            ['\u6210\u4ea4\u91cf MA(20)', '\u6210\u4ea4\u91cf\u5747\u7ebf', '\u786e\u8ba4\u8d8b\u52bf'],
            ['\u6536\u76d8\u4ef7 vs 24h \u524d', '\u77ed\u671f\u6536\u76ca\u7387', '\u52a8\u91cf'],
        ]
    )

    p = doc.add_paragraph()
    run = p.add_run('\u98ce\u9669\u63d0\u793a\uff1a')
    run.bold = True
    run = p.add_run('Kronos \u9884\u8bad\u7ec3\u53ea\u63a5\u53d7 OHLCV\uff0c\u989d\u5916\u7279\u5f81\u9700\u8981\u5fae\u8c03\u6216\u6362\u7528\u5176\u4ed6\u6a21\u578b')
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

    # --- Plan 5 ---
    doc.add_heading('\u65b9\u6848\u4e94\uff1a\u96c6\u6210\u9884\u6d4b\uff08Ensemble\uff09', level=2)
    p = doc.add_paragraph()
    run = p.add_run('\u95ee\u9898\uff1a')
    run.bold = True
    p.add_run('\u5355\u6a21\u578b\u9884\u6d4b\u65b9\u5dee\u5927')

    p = doc.add_paragraph()
    run = p.add_run('\u884c\u52a8\u8ba1\u5212\uff1a')
    run.bold = True
    doc.add_paragraph('\u7528\u4e0d\u540c random seed \u8dd1\u591a\u6b21\u9884\u6d4b\uff08Temperature sampling\uff09', style='List Bullet')
    doc.add_paragraph('\u5bf9\u9884\u6d4b\u7ed3\u679c\u53d6 median\uff08\u6bd4 mean \u66f4\u6297\u5f02\u5e38\u503c\uff09', style='List Bullet')
    doc.add_paragraph('\u6216\u8005\uff1amini + small + base \u4e09\u4e2a\u6a21\u578b\u9884\u6d4b\u53d6\u52a0\u6743\u5e73\u5747', style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('\u9884\u671f\u6548\u679c\uff1a')
    run.bold = True
    p.add_run('\u65b9\u5411\u51c6\u786e\u7387\u4ece 69% \u2192 75~80%')

    # --- Plan 6 ---
    doc.add_heading('\u65b9\u6848\u516d\uff1a\u6df7\u5408\u6a21\u578b\u67b6\u6784\uff08\u6839\u672c\u89e3\u51b3\u65b9\u6848\uff09', level=2)
    p = doc.add_paragraph()
    run = p.add_run('\u95ee\u9898\uff1a')
    run.bold = True
    p.add_run('Kronos \u662f\u901a\u7528\u65f6\u95f4\u5e8f\u6a21\u578b\uff0c\u4e0d\u662f\u4e13\u4e3a\u91d1\u878d\u9884\u6d4b\u8bbe\u8ba1\u7684')

    add_table(doc,
        ['\u6a21\u578b', '\u7c7b\u578b', '\u4f18\u70b9', '\u7f3a\u70b9'],
        [
            ['TFT\uff08Temporal Fusion Transformer\uff09', '\u4e13\u7528\u91d1\u878d\u6a21\u578b', '\u5185\u7f6e\u7279\u5f81\u91cd\u8981\u6027', '\u9700\u8981\u8bad\u7ec3'],
            ['N-BEATS', '\u65f6\u5e8f\u4e13\u7528', 'M4 \u7ade\u8d5b\u8868\u73b0\u4f18\u5f02', '\u9700\u8981\u8bad\u7ec3'],
            ['PatchTST', 'Transformer \u65f6\u5e8f', '\u6700\u65b0 SOTA', '\u9700\u8981\u8bad\u7ec3'],
            ['Kronos + \u7ebf\u6027\u56de\u5f52\u6821\u51c6', '\u6df7\u5408\u65b9\u6848', '\u5feb\u901f\u90e8\u7f72', '\u6548\u679c\u6709\u9650'],
        ]
    )

    p = doc.add_paragraph()
    run = p.add_run('\u63a8\u8350\u6df7\u5408\u65b9\u6848\uff1a')
    run.bold = True
    run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
    doc.add_paragraph('\u4fdd\u7559 Kronos \u4f5c\u4e3a\u201c\u7279\u5f81\u63d0\u53d6\u5668\u201d', style='List Bullet')
    doc.add_paragraph('\u7528 Kronos \u7684\u8f93\u51fa + \u6280\u672f\u6307\u6807\u4f5c\u4e3a\u7279\u5f81', style='List Bullet')
    doc.add_paragraph('\u4e0a\u5c42\u52a0\u4e00\u4e2a\u8f7b\u91cf\u7ebf\u6027\u56de\u5f52 / XGBoost \u505a\u6700\u7ec8\u9884\u6d4b', style='List Bullet')
    doc.add_paragraph('\u4e0d\u9700\u8981\u91cd\u65b0\u8bad\u7ec3 Kronos\uff0c\u53ea\u9700\u8981\u8bad\u7ec3\u4e0a\u5c42\u6a21\u578b\uff08< 1000 \u6837\u672c\u5373\u53ef\uff09', style='List Bullet')

    # ===== Section 3: Execution Roadmap =====
    doc.add_heading('\u4e09\u3001\u63a8\u8350\u6267\u884c\u987a\u5e8f', level=1)

    add_table(doc,
        ['\u6b65\u9aa4', '\u65f6\u95f4', '\u65b9\u6848', '\u5185\u5bb9'],
        [
            ['\u7b2c 1 \u6b65', '\u4eca\u5929', '\u65b9\u6848\u4e00', '\u8c03 Temperature + \u540e\u6821\u51c6'],
            ['\u7b2c 2 \u6b65', '\u672c\u5468', '\u65b9\u6848\u4e8c', '\u4e0b\u8f7d small/base \u6a21\u578b\u5bf9\u6bd4'],
            ['\u7b2c 3 \u6b65', '\u672c\u5468', '\u65b9\u6848\u4e09', '\u5b8c\u6574\u56de\u6d4b\uff0c\u786e\u8ba4\u6a21\u578b\u771f\u5b9e\u80fd\u529b'],
            ['\u7b2c 4 \u6b65', '\u4e0b\u5468', '\u65b9\u6848\u516d', '\u6df7\u5408\u6a21\u578b\uff0c\u4e0a\u7ebf\u66f4\u53ef\u9760\u7684\u9884\u6d4b'],
        ]
    )

    # ===== Section 4: Code Implementation =====
    doc.add_heading('\u56db\u3001\u5feb\u901f\u89c1\u6548\uff1a\u540e\u6821\u51c6\u5b9e\u73b0', level=1)

    p = doc.add_paragraph('\u5728 kronos_dashboard.py \u4e2d\u52a0\u5165\u6821\u51c6\u529f\u80fd\u7684\u5173\u952e\u4ee3\u7801\uff1a')

    code = doc.add_paragraph()
    code.style = 'Normal'
    code_text = (
        'def calibrate_predictions(pred_df, actual_df, method="scalar"):\n'
        '    """\n'
        '    \u7528\u5386\u53f2\u5b9e\u9645\u6570\u636e\u6821\u51c6\u9884\u6d4b\u7ed3\u679c\n'
        '    method: "scalar" \u6807\u91cf\u7f29\u653e, "quantile" \u5206\u4f4d\u6570\u6620\u5c04\n'
        '    """\n'
        '    if method == "scalar":\n'
        '        pred_mean = pred_df[\'close\'].mean()\n'
        '        actual_mean = actual_df[\'close\'].mean()\n'
        '        factor = actual_mean / pred_mean\n'
        '        for col in [\'open\', \'high\', \'low\', \'close\']:\n'
        '            pred_df[col] = pred_df[col] * factor\n'
        '        return pred_df, factor'
    )
    run = code.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run('\u52a0\u5165\u540e\uff0c\u5728\u9884\u6d4b\u5b8c\u6210\u540e\u81ea\u52a8\u663e\u793a\u201c\u6821\u51c6\u540e MAPE\u201d\uff0c\u9884\u8ba1\u4ece 30% \u964d\u81f3 15% \u5de6\u53f3\u3002')

    # ===== Section 5: Success Metrics =====
    doc.add_heading('\u4e94\u3001\u6210\u529f\u6307\u6807\uff08\u4f18\u5316\u76ee\u6807\uff09', level=1)

    add_table(doc,
        ['\u6307\u6807', '\u5f53\u524d\u503c', '\u76ee\u6807\u503c', '\u4f18\u5316\u7a7a\u95f4'],
        [
            ['MAPE', '30.17%', '< 15%', '\u964d\u4f4e 50%+'],
            ['\u65b9\u5411\u51c6\u786e\u7387', '69%', '> 75%', '\u63d0\u5347 6%+'],
            ['\u7cfb\u7edf\u6027\u504f\u5dee', '+$21K\uff08\u504f\u9ad8\uff09', '< \u00b1$5K', '\u964d\u4f4e 76%+'],
            ['\u9884\u6d4b\u533a\u95f4\u8986\u76d6\u7387', '\u672a\u77e5', '80~90%', '\u65b0\u589e\u6307\u6807'],
        ],
        header_color="27AE60"
    )

    # Footer
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2014 Kronos \u91d1\u878d\u9884\u6d4b\u5de5\u4f5c\u53f0 \u00b7 \u6a21\u578b\u4f18\u5316\u8ba1\u5212 \u2014')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # Save
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "kronos_optimization_plan.docx")
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_doc()
    print("Done!")
